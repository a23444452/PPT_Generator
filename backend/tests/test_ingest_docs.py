"""docx / pdf converter 測試。fixtures 動態產生（不依賴外部檔案）。"""

import docx
import fitz
import pytest
from PIL import Image

from app.ingest import IngestError, ingest_file
from app.store.project import create_project


@pytest.fixture
def project(tmp_path):
    return create_project(tmp_path / "projects", "docs 測試")


# ---------- docx ----------


@pytest.fixture
def sample_docx(tmp_path):
    """含 Heading 1、內文段落、表格的 docx fixture。"""
    d = docx.Document()
    d.add_heading("標題一", level=1)
    d.add_paragraph("這是內文段落。")
    d.add_heading("子標題", level=2)
    table = d.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "項目"
    table.cell(0, 1).text = "數量"
    table.cell(1, 0).text = "蘋果"
    table.cell(1, 1).text = "42"
    path = tmp_path / "report.docx"
    d.save(path)
    return path


def test_docx_heading1_becomes_h1(project, sample_docx):
    result = ingest_file(sample_docx, project)

    assert result.output_type == "markdown"
    assert result.output_path == project.path / "md" / "report.docx.md"
    md = result.output_path.read_text(encoding="utf-8")

    assert "# 標題一" in md
    assert "## 子標題" in md
    assert "這是內文段落。" in md


def test_docx_table_becomes_markdown_table(project, sample_docx):
    result = ingest_file(sample_docx, project)
    md = result.output_path.read_text(encoding="utf-8")

    assert "| 項目 | 數量 |" in md
    assert "| 蘋果 | 42 |" in md
    # 需有分隔線
    assert "| --- | --- |" in md


def test_docx_body_order_preserved(project, sample_docx):
    """段落與表格依原始順序輸出（標題在表格之前）。"""
    result = ingest_file(sample_docx, project)
    md = result.output_path.read_text(encoding="utf-8")

    heading_pos = md.index("# 標題一")
    table_pos = md.index("| 項目 | 數量 |")
    assert heading_pos < table_pos


def test_corrupt_docx_raises_friendly_ingest_error(tmp_path, project):
    src = tmp_path / "broken.docx"
    src.write_bytes(b"this is not a zip archive")
    with pytest.raises(IngestError) as excinfo:
        ingest_file(src, project)
    msg = str(excinfo.value)
    assert "broken.docx" in msg
    assert "Traceback" not in msg


# ---------- pdf ----------


@pytest.fixture
def sample_pdf(tmp_path):
    """單頁 PDF：大字級標題 + 一般字級內文。"""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Big Title", fontsize=28)
    page.insert_text((72, 150), "This is body text content here.", fontsize=12)
    page.insert_text((72, 180), "More body text follows here.", fontsize=12)
    path = tmp_path / "paper.pdf"
    doc.save(path)
    doc.close()
    return path


@pytest.fixture
def sample_pdf_with_images(tmp_path):
    """含一張小圖（應跳過）與一張大圖（應抽出）的 PDF，且大圖重複貼兩次。"""
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Image Page", fontsize=24)

    tiny_path = tmp_path / "tiny_src.png"
    Image.new("RGB", (50, 50), "red").save(tiny_path)
    big_path = tmp_path / "big_src.png"
    Image.new("RGB", (200, 200), "blue").save(big_path)

    page.insert_image(fitz.Rect(0, 200, 60, 260), filename=str(tiny_path))
    page.insert_image(fitz.Rect(100, 300, 300, 500), filename=str(big_path))
    # 同一張大圖再貼一次到第二頁，測試 MD5 去重
    page2 = doc.new_page()
    page2.insert_image(fitz.Rect(100, 300, 300, 500), filename=str(big_path))

    path = tmp_path / "withimg.pdf"
    doc.save(path)
    doc.close()
    return path


def test_pdf_largest_font_line_becomes_heading(project, sample_pdf):
    result = ingest_file(sample_pdf, project)

    assert result.output_type == "markdown"
    assert result.output_path == project.path / "md" / "paper.pdf.md"
    md = result.output_path.read_text(encoding="utf-8")

    assert "# Big Title" in md


def test_pdf_body_text_preserved(project, sample_pdf):
    result = ingest_file(sample_pdf, project)
    md = result.output_path.read_text(encoding="utf-8")

    assert "This is body text content here." in md
    assert "More body text follows here." in md


def test_pdf_notes_table_may_be_missing(project, sample_pdf):
    result = ingest_file(sample_pdf, project)
    md = result.output_path.read_text(encoding="utf-8")

    assert "PDF 表格可能遺漏" in md
    assert any("表格" in w for w in result.warnings)


def test_pdf_no_images_has_no_extra_assets(project, sample_pdf):
    result = ingest_file(sample_pdf, project)
    assert result.extra_assets == ()


def test_pdf_images_extracted_to_assets(project, sample_pdf_with_images):
    result = ingest_file(sample_pdf_with_images, project)
    md = result.output_path.read_text(encoding="utf-8")

    # 小圖（< 100x100）被跳過，只剩一張大圖（去重後）
    assert len(result.extra_assets) == 1
    asset_path = result.extra_assets[0]
    assert asset_path.parent == project.path / "assets"
    assert asset_path.exists()

    assert f"![](assets/{asset_path.name})" in md


def test_pdf_small_image_skipped(project, sample_pdf_with_images):
    result = ingest_file(sample_pdf_with_images, project)
    # 只有大圖被抽出，確認檔名不對應到 tiny
    for asset in result.extra_assets:
        img = Image.open(asset)
        assert img.width >= 100
        assert img.height >= 100


def test_pdf_duplicate_image_deduplicated_by_md5(project, sample_pdf_with_images):
    result = ingest_file(sample_pdf_with_images, project)
    # 同一張圖在兩頁各貼一次，理應只保留一份
    assert len(result.extra_assets) == 1


def test_corrupt_pdf_raises_friendly_ingest_error(tmp_path, project):
    src = tmp_path / "broken.pdf"
    src.write_bytes(b"this is not a pdf file")
    with pytest.raises(IngestError) as excinfo:
        ingest_file(src, project)
    msg = str(excinfo.value)
    assert "broken.pdf" in msg
    assert "Traceback" not in msg
